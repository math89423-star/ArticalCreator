import re
import io
import os
import matplotlib.pyplot as plt
import seaborn as sns
import base64
import pandas as pd  
import numpy as np
from matplotlib import font_manager
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def register_custom_font():
    """
    自动查找并加载中文字体
    查找顺序：本地文件 -> Linux常见字体 -> Windows常见字体
    """
    current_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(current_dir)
    
    possible_font_paths = [
        'SimHei.ttf',
        os.path.join(current_dir, 'SimHei.ttf'),
        os.path.join(project_root, 'SimHei.ttf'),
        '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
        '/usr/share/fonts/truetype/arphic/ukai.ttc'
    ]

    for font_path in possible_font_paths:
        if os.path.exists(font_path):
            try:
                prop = font_manager.FontProperties(fname=font_path)
                font_manager.fontManager.addfont(font_path)
                # print(f"[Font] 成功加载字体文件: {font_path}")
                plt.rcParams['font.sans-serif'] = [prop.get_name()]
                plt.rcParams['axes.unicode_minus'] = False
                return prop.get_name()
            except Exception as e:
                print(f"[Font] 尝试加载 {font_path} 失败: {e}")

    fonts_to_try = [
        'WenQuanYi Micro Hei', 'WenQuanYi Zen Hei', 'Noto Sans CJK SC',
        'Droid Sans Fallback', 'SimHei', 'Microsoft YaHei', 'SimSun'
    ]
    
    system_fonts = {f.name for f in font_manager.fontManager.ttflist}
    
    for f in fonts_to_try:
        if f in system_fonts:
            plt.rcParams['font.sans-serif'] = [f]
            plt.rcParams['axes.unicode_minus'] = False
            return f

    print("[Font] ❌ 警告: 未找到任何中文字体，中文将无法显示！请上传 SimHei.ttf 到项目目录。")
    return 'sans-serif'

# 初始化字体
CURRENT_FONT_NAME = register_custom_font()

class TextCleaner:
    @staticmethod
    def clean_special_chars(text):
        # 1. 统一换行符
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        # 2. 替换全角空格和其他不可见字符为普通空格
        text = text.replace('\u3000', ' ').replace('\u00A0', ' ').replace('\u200b', '')
        return text

    @staticmethod
    def fix_table_newlines(text):
        """
        【核心修复】解决 "文字 | 表头 |" 同行粘连问题。
        如果不修复，Pandoc或Docx解析会将表格当做普通文本。
        """
        if not text: return ""
        
        # 先清洗干扰字符
        text = TextCleaner.clean_special_chars(text)
        
        lines = text.split('\n')
        final_lines = []
        
        for line in lines:
            stripped = line.strip()
            if not stripped:
                final_lines.append(line)
                continue
            
            # 检测是否包含表格特征：
            # 1. 包含竖线
            # 2. 竖线不是第一个字符（说明前面有文字粘连）
            # 3. 包含至少两个竖线（构成表格结构）
            pipe_count = stripped.count('|')
            first_pipe_idx = stripped.find('|')
            
            if pipe_count >= 2 and first_pipe_idx > 0:
                pre_content = stripped[:first_pipe_idx]
                table_content = stripped[first_pipe_idx:]
                
                # 如果前面有实义文本（不仅仅是列表符 - 或 * 或 1.）
                if pre_content.strip() and not pre_content.strip() in ['-', '*', '1.']:
                    # 命中粘连！拆分它
                    final_lines.append(pre_content.strip())
                    final_lines.append('') # 强制插入空行
                    final_lines.append(table_content)
                    continue

            # 普通行处理：如果是正常的表格行（以 | 开头），确保它上方有空行
            if stripped.startswith('|') and pipe_count >= 1:
                if final_lines:
                    prev = final_lines[-1].strip()
                    # 上一行有内容，且不是表格 -> 插入空行
                    if prev and not prev.startswith('|'):
                        final_lines.append('')
            
            final_lines.append(line)
            
        return '\n'.join(final_lines)
    
    @staticmethod
    def correct_punctuation(text):
        masks = {}
        def save_mask(match):
            key = f"__MASK_{len(masks)}__"
            masks[key] = match.group(0)
            return key

        text = re.sub(r'```[\s\S]*?```', save_mask, text)
        text = re.sub(r'`[^`\n]+`', save_mask, text)
        text = re.sub(r'!\[.*?\]\(.*?\)', save_mask, text)
        text = re.sub(r'\[.*?\]\(.*?\)', save_mask, text)
        
        # 智能标点替换
        pairs = [
            (r'(?<=[\u4e00-\u9fa5]),', '，'),
            (r',(?=[\u4e00-\u9fa5])', '，'),
            (r'(?<=[\u4e00-\u9fa5])\.', '。'),
            (r'(?<=[\u4e00-\u9fa5]):', '：'),
            (r':(?=[\u4e00-\u9fa5])', '：'),
            (r'(?<=[\u4e00-\u9fa5]);', '；'),
            (r';(?=[\u4e00-\u9fa5])', '；'),
            (r'(?<=[\u4e00-\u9fa5])\?', '？'),
            (r'\?(?=[\u4e00-\u9fa5])', '？'),
            (r'(?<=[\u4e00-\u9fa5])!', '！'),
            (r'!(?=[\u4e00-\u9fa5])', '！'),
            (r'(?<=[\u4e00-\u9fa5])\s*\(', '（'),
            (r'\((?=[\u4e00-\u9fa5])', '（'),
            (r'(?<=[\u4e00-\u9fa5])\)', '）'),
            (r'\)(?=[\u4e00-\u9fa5])', '）')
        ]
        
        for p, r in pairs:
            text = re.sub(p, r, text)

        def quote_replacer(match):
            content = match.group(1)
            if re.search(r'[\u4e00-\u9fa5]', content):
                return f'“{content}”'
            return match.group(0)
        text = re.sub(r'"(.*?)"', quote_replacer, text, flags=re.DOTALL)

        for key in reversed(list(masks.keys())):
            text = text.replace(key, masks[key])
        return text
    
    @staticmethod
    def convert_cn_numbers(text: str) -> str:
        return text

class MarkdownToDocx:
    @staticmethod
    def set_table_borders(table):
        """
        设置学术三线表样式：顶底粗线，中间细线，无竖线
        """
        tbl = table._tbl
        tblPr = tbl.tblPr
        # 清除默认样式
        for tag in ['w:tblStyle', 'w:tblBorders', 'w:tblLook']:
            element = tblPr.find(qn(tag))
            if element is not None: tblPr.remove(element)
            
        tblBorders = OxmlElement('w:tblBorders')
        
        def border(tag, val, sz, space="0", color="auto"):
            el = OxmlElement(f'w:{tag}')
            el.set(qn('w:val'), val)
            el.set(qn('w:sz'), str(sz)) # 1/8 pt. 24=3pt(显示为1.5磅), 4-6=细线
            el.set(qn('w:space'), space)
            el.set(qn('w:color'), color)
            return el
            
        # 顶线 (加粗)
        tblBorders.append(border('top', 'single', 24))
        # 底线 (加粗)
        tblBorders.append(border('bottom', 'single', 24))
        # 内部横线 (细线)
        tblBorders.append(border('insideH', 'single', 6))
        # 去除竖线
        tblBorders.append(border('left', 'nil', 0))
        tblBorders.append(border('right', 'nil', 0))
        tblBorders.append(border('insideV', 'nil', 0))
        
        tblPr.append(tblBorders)

    @staticmethod
    def create_error_image(error_msg):
        plt.close('all')
        plt.figure(figsize=(8, 4))
        plt.text(0.5, 0.5, f"Plot Error:\n{error_msg}", 
                 horizontalalignment='center', verticalalignment='center',
                 fontsize=12, color='red', bbox=dict(facecolor='#ffe6e6', edgecolor='red'))
        plt.axis('off')
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=100)
        plt.close('all')
        buf.seek(0)
        return buf

    @staticmethod
    def exec_python_plot(code_str):
        try:
            code_str = re.sub(r'^```python', '', code_str.strip(), flags=re.MULTILINE|re.IGNORECASE)
            code_str = re.sub(r'^```', '', code_str.strip(), flags=re.MULTILINE)
            code_str = code_str.replace('\u3000', ' ').replace('\u00A0', ' ').replace('\u200b', '')
            code_str = re.sub(r"plt\.rcParams\[.*?\]\s*=\s*.*", "", code_str)
            code_str = re.sub(r"sns\.set_theme\(.*?\)", "", code_str) 
            code_str = re.sub(r"sns\.set\(.*?\)", "", code_str)
            plt.close('all') 
            plt.clf()
            sns.set_theme(style="whitegrid")
            plt.rcParams['font.sans-serif'] = [CURRENT_FONT_NAME]
            plt.rcParams['axes.unicode_minus'] = False
            local_vars = {'plt': plt, 'sns': sns, 'pd': pd, 'np': np}
            exec(code_str, {}, local_vars)
            buf = io.BytesIO()
            plt.tight_layout()
            plt.savefig(buf, format='png', dpi=300, bbox_inches='tight') 
            plt.close('all')
            buf.seek(0)
            return buf
        except Exception as e:
            print(f"Python Plot Error: {e}")
            return MarkdownToDocx.create_error_image(str(e))

    @staticmethod
    def parse_markdown_table(lines, start_idx):
        i = start_idx
        header_line = lines[i].strip()
        if not header_line.startswith('|'): return None, i
        headers = [c.strip() for c in header_line.strip('|').split('|')]
        i += 1
        # 兼容中文破折号和冒号
        separator_line = lines[i].strip() if i < len(lines) else ""
        if i < len(lines) and re.match(r'^[|\-\s:—–]+$', separator_line): i += 1
        else: return None, start_idx
            
        data = [headers]
        while i < len(lines):
            line = lines[i].strip()
            if not line.startswith('|'): break
            row = [c.strip() for c in line.strip('|').split('|')]
            if len(row) < len(headers): row += [''] * (len(headers) - len(row))
            else: row = row[:len(headers)]
            data.append(row)
            i += 1
        return data, i

    @staticmethod
    def convert(markdown_text):
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = u'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        
        # 1. 【关键】先修复表格粘连问题
        markdown_text = TextCleaner.fix_table_newlines(markdown_text)
        # 2. 标点修正
        markdown_text = TextCleaner.correct_punctuation(markdown_text)
        
        lines = markdown_text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line: 
                i += 1
                continue
            
            # 标题
            if line.startswith('#'):
                level = len(line.split(' ')[0])
                content = line.lstrip('#').strip()
                content = re.sub(r'^[•\*\-\s]+', '', content)
                doc_level = min(level, 3)
                heading = doc.add_heading('', level=doc_level)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
                run = heading.add_run(content)
                run.font.name = u'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
                run.font.color.rgb = RGBColor(0, 0, 0)
                heading.paragraph_format.keep_with_next = False
                if level == 1:
                    run.font.size = Pt(16)
                    heading.paragraph_format.space_before = Pt(12)
                    heading.paragraph_format.space_after = Pt(12)
                elif level == 2:
                    run.font.size = Pt(14)
                    heading.paragraph_format.space_before = Pt(12)
                    heading.paragraph_format.space_after = Pt(12)
                else:
                    run.font.size = Pt(13)
                    heading.paragraph_format.space_before = Pt(6)
                    heading.paragraph_format.space_after = Pt(6)
                i += 1
                continue
            
            # Base64 图片
            img_match = re.search(r'!\[.*?\]\(data:image\/png;base64,(.*?)\)', line)
            if not img_match:
                img_match = re.search(r'src=["\']data:image\/png;base64,(.*?)["\']', line)
            if img_match:
                try:
                    base64_str = img_match.group(1)
                    img_data = base64.b64decode(base64_str)
                    img_stream = io.BytesIO(img_data)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(img_stream, width=Cm(14)) 
                except Exception as e:
                    print(f"Base64 Image Error: {e}")
                i += 1
                continue

            # Python 代码块
            if line.startswith('```python'):
                code_block = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_block.append(lines[i])
                    i += 1
                i += 1
                img_stream = MarkdownToDocx.exec_python_plot("\n".join(code_block))
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                if img_stream:
                    run.add_picture(img_stream, width=Cm(14))
                continue
            
            # 表格处理
            if line.startswith('|'):
                table_data, next_i = MarkdownToDocx.parse_markdown_table(lines, i)
                if table_data and len(table_data) > 1:
                    rows = len(table_data)
                    cols = max(len(r) for r in table_data)
                    if cols > 0:
                        table = doc.add_table(rows=rows, cols=cols)
                        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        table.autofit = True
                        for r, row_data in enumerate(table_data):
                            for c, cell_text in enumerate(row_data):
                                if c >= cols: break
                                cell = table.cell(r, c)
                                cell.text = ""
                                p = cell.paragraphs[0]
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                clean_text = cell_text.replace('**', '').strip()
                                clean_text = TextCleaner.correct_punctuation(clean_text)
                                run = p.add_run(clean_text)
                                run.font.name = u'Times New Roman'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                                run.font.size = Pt(10.5)
                                if r == 0: run.bold = True
                        try: 
                            # 应用三线表
                            MarkdownToDocx.set_table_borders(table)
                        except Exception as e: 
                            print(f"Table formatting error: {e}")
                        
                        doc.add_paragraph()
                        i = next_i
                        continue
            
            # 正文段落
            if line:
                clean_line = line.replace('**', '').strip() 
                line = TextCleaner.correct_punctuation(line)
                is_caption = re.match(r'^(图|表)\s*\d+[\.\-]\d+', clean_line) or re.match(r'^(图|表)\s*\d+', clean_line)
                p = doc.add_paragraph()
                if is_caption:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.first_line_indent = None
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.first_line_indent = Pt(24) 
                    p.paragraph_format.line_spacing = 1.5
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if not part: continue
                    is_bold = part.startswith('**') and part.endswith('**')
                    text_content = part.replace('**', '') if is_bold else part
                    run = p.add_run(text_content)
                    run.font.name = u'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                    run.font.size = Pt(12)
                    run.bold = is_bold
            i += 1
        out_stream = io.BytesIO()
        doc.save(out_stream)
        out_stream.seek(0)
        return out_stream
    
class TextReportParser:
    @staticmethod
    def parse(text: str) -> dict:
        """
        解析开题报告文本
        """
        data = { "title": "", "outline_content": "", "cn_refs": [], "en_refs": [], "review": "" }
        if not text: return data
        title_match = re.search(r'(?:论文)?(?:题目|Title)[:：]\s*(.*)', text, re.IGNORECASE)
        if title_match: data["title"] = title_match.group(1).strip()
        else:
            lines = [l.strip() for l in text.split('\n') if l.strip()]
            if lines and len(lines[0]) < 100: data["title"] = lines[0]
        
        ref_split = re.split(r'^\s*(?:参考文献|References)[:：]?\s*$', text, flags=re.MULTILINE | re.IGNORECASE)
        main_body = text 
        if len(ref_split) > 1:
            ref_text = ref_split[-1].strip()
            main_body = ref_split[0] 
            refs = [line.strip() for line in ref_text.split('\n') if line.strip()]
            for ref in refs:
                clean_ref = re.sub(r'^(\[\d+\]|\d+\.|（\d+）|\(\d+\))\s*', '', ref)
                if len(clean_ref) < 5: continue
                if re.search(r'[\u4e00-\u9fa5]', clean_ref): data["cn_refs"].append(clean_ref)
                else: data["en_refs"].append(clean_ref)
        
        review_keywords = ["文献综述", "研究现状", "国内外研究", "Literature Review", "Related Work"]
        for kw in review_keywords:
            pattern = rf'{kw}[:：]?\s*([\s\S]*?)(?=(?:研究内容|研究方法|论文提纲|论文目录|第[一二三]章|3\.|^3\s|Chapter|Methodology|Research Content)|$)'
            match = re.search(pattern, main_body, re.IGNORECASE | re.MULTILINE)
            if match:
                review_content = match.group(1).strip()
                if len(review_content) > 50:
                    data["review"] = review_content
                    break
        
        outline_match = re.search(r'(?:目录|提纲|章节安排|结构安排|Table of Contents|Outline)[:：]?\s*([\s\S]*)', main_body, re.MULTILINE | re.IGNORECASE)
        source_for_outline = outline_match.group(1) if outline_match else main_body
        outline_lines = []
        raw_lines = source_for_outline.split('\n')
        outline_patterns = [
            r'^Chapter\s+\d+', r'^Section\s+\d+', r'^Part\s+\d+', 
            r'^第[一二三四五六七八九十0-9]+章', r'^\d+(\.\d+)*\s', r'^\d+\.\s', 
            r'^[一二三四五六]+、', 
            r'^(?:Abstract|Introduction|Literature Review|Methodology|Results|Discussion|Conclusion|References|Appendix)',
            r'^(?:摘要|绪论|结论|参考文献|致谢|附录)'
        ]
        combined_pattern = '|'.join(outline_patterns)
        for line in raw_lines:
            line = line.strip()
            if len(line) < 100 and re.match(combined_pattern, line, re.IGNORECASE): 
                outline_lines.append(line)
        if outline_lines: data["outline_content"] = "\n".join(outline_lines)
        return data