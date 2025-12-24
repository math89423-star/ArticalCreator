import json
import os
import secrets
import io
import threading
import time
from collections import defaultdict

import matplotlib
# 设置后端为 Agg，确保在无显示器的服务器环境下也能运行
matplotlib.use('Agg') 

# [新增依赖库]
import pandas as pd  # 处理 Excel/CSV
import pypdf         # 处理 PDF
import docx          # 处理 Word .docx

# 引入 waitress (请确保 pip install waitress)
from waitress import serve
from flask import Flask, render_template, request, Response, stream_with_context, jsonify, send_file, session
from utils.word import MarkdownToDocx
from utils.prompts import PaperAutoWriter

app = Flask(__name__)
app.secret_key = "super_secret_key_for_session" # 用于管理员登录Session

# ==============================================================================
# 配置区域
# ==============================================================================
API_KEY = "sk-VuSl3xg7XTQUbWzs4QCeinJk70H4rhFUtrLdZlBC6hvjvs1t" 
BASE_URL = "https://tb.api.mkeai.com/v1"
MODEL_NAME = "deepseek-v3.2" 

# 管理员账号配置
ADMIN_USERNAME = "admin"
ADMIN_PASSWORD = "admin123"

# 卡密存储文件
KEYS_FILE = "valid_keys.json"

# ==============================================================================
# 数据持久化与鉴权逻辑
# ==============================================================================
def load_keys():
    if not os.path.exists(KEYS_FILE):
        return []
    try:
        with open(KEYS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except:
        return []

def save_keys(keys):
    with open(KEYS_FILE, 'w', encoding='utf-8') as f:
        json.dump(keys, f)

# 内存中缓存有效卡密
VALID_KEYS = set(load_keys())

def is_valid_key(key):
    return key in VALID_KEYS

# --- 任务管理器 (优化版) ---
class TaskManager:
    def __init__(self):
        # 改用 RLock (可重入锁)，更加安全，防止自身死锁
        self._lock = threading.RLock()
        self._user_tasks = defaultdict(dict)

    def start_task(self, user_id, task_id):
        """初始化任务状态"""
        with self._lock:
            self._user_tasks[user_id][task_id] = {
                'status': 'running',
                'events': [],      # 消息缓存队列
                'created_at': time.time(),
                'last_read_idx': 0 
            }

    def append_event(self, user_id, task_id, event_data):
        """后台线程写入消息"""
        with self._lock:
            if task_id in self._user_tasks[user_id]:
                self._user_tasks[user_id][task_id]['events'].append(event_data)

    def get_events_from(self, user_id, task_id, start_index):
        """前端读取消息（增量读取）"""
        with self._lock:
            task = self._user_tasks[user_id].get(task_id)
            if not task:
                return [], 'stopped'
            
            # 安全获取切片，即使 index 越界也不会报错
            events_len = len(task['events'])
            if start_index >= events_len:
                return [], task['status']
                
            new_events = task['events'][start_index:]
            return new_events, task['status']

    def set_status(self, user_id, task_id, status):
        with self._lock:
            if task_id in self._user_tasks[user_id]:
                self._user_tasks[user_id][task_id]['status'] = status

    def get_status(self, user_id, task_id):
        with self._lock:
            return self._user_tasks[user_id].get(task_id, {}).get('status', 'stopped')

task_manager = TaskManager()

# ==============================================================================
# 多格式文件内容提取工具
# ==============================================================================
def extract_file_content(file_stream, filename) -> str:
    """
    根据文件后缀名，提取文件内容为纯文本字符串。
    注意：file_stream 必须是 BytesIO 或已打开的文件对象
    """
    filename = filename.lower()
    content = ""
    
    try:
        # 重置指针，防止读取位置错误
        if hasattr(file_stream, 'seek'):
            file_stream.seek(0)

        # 1. Excel/CSV
        if filename.endswith('.csv'):
            try:
                df = pd.read_csv(file_stream)
            except UnicodeDecodeError:
                file_stream.seek(0)
                df = pd.read_csv(file_stream, encoding='gbk')
            content = f"\n【文件 {filename} 数据预览(前60行)】:\n" + df.head(60).to_markdown(index=False)
        
        elif filename.endswith(('.xls', '.xlsx')):
            df = pd.read_excel(file_stream)
            content = f"\n【文件 {filename} 数据预览(前60行)】:\n" + df.head(60).to_markdown(index=False)
            
        # 2. TXT
        elif filename.endswith('.txt'):
            content = f"\n【文件 {filename} 内容】:\n"
            try:
                text = file_stream.read().decode('utf-8')
            except:
                file_stream.seek(0)
                text = file_stream.read().decode('gbk', errors='ignore')
            content += text[:5000]
            
        # 3. PDF
        elif filename.endswith('.pdf'):
            reader = pypdf.PdfReader(file_stream)
            text = ""
            for i, page in enumerate(reader.pages[:15]): 
                page_text = page.extract_text()
                if page_text: 
                    text += f"[第{i+1}页] {page_text}\n"
            content = f"\n【文件 {filename} 内容提取】:\n{text}"

        # 4. DOCX
        elif filename.endswith('.docx'):
            doc = docx.Document(file_stream)
            text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    text += " | ".join(row_text) + "\n"
            content = f"\n【文件 {filename} 内容提取】:\n{text[:5000]}"

        else:
            content = f"\n【文件 {filename}】: 暂不支持该格式解析。"
            
    except Exception as e:
        print(f"解析文件 {filename} 失败: {e}")
        content = f"\n【文件 {filename}】: 解析失败 - {str(e)}"
        
    return content

# 1. 后台工作线程函数 (现在负责所有重活)
def background_worker(writer, task_id, title, chapters, references, text_custom_data, raw_files_data, check_status_func, initial_context, user_id):
    try:
        # 1. 在后台线程中进行文件解析 (耗时操作放在这里，不阻塞主线程)
        final_custom_data = text_custom_data
        
        if raw_files_data:
            # 发送日志告诉前端正在解析文件
            task_manager.append_event(user_id, task_id, f"data: {json.dumps({'type': 'log', 'msg': '📂 正在后台解析上传的文件...'})}\n\n")
            
            file_extracted_text = ""
            for file_info in raw_files_data:
                # 🚀 关键：每解析一个文件，主动休眠 10ms 释放 GIL 锁，防止卡死其他正在生成的任务
                time.sleep(0.01) 
                
                try:
                    extracted = extract_file_content(file_info['content'], file_info['name'])
                    file_extracted_text += extracted + "\n\n"
                except Exception as e:
                    file_extracted_text += f"\n文件 {file_info['name']} 解析失败: {e}\n"
            
            final_custom_data = text_custom_data + "\n" + file_extracted_text
            task_manager.append_event(user_id, task_id, f"data: {json.dumps({'type': 'log', 'msg': '✅ 文件解析完成，开始生成...'})}\n\n")

        # 2. 执行生成器
        generator = writer.generate_stream(
            task_id, title, chapters, references, final_custom_data, check_status_func, initial_context
        )
        
        # 3. 逐条消费
        for chunk in generator:
            task_manager.append_event(user_id, task_id, chunk)
            # 极短暂休眠，释放GIL锁，让其他并发任务的SSE线程有机会呼吸
            time.sleep(0.005) 
            
    except Exception as e:
        error_msg = json.dumps({'type': 'log', 'msg': f'❌ 后台任务异常: {str(e)}'})
        task_manager.append_event(user_id, task_id, f"data: {error_msg}\n\n")
    finally:
        # 无论成功还是失败，都要确保将状态标记为完成或停止
        current_status = task_manager.get_status(user_id, task_id)
        if current_status == 'running':
            task_manager.set_status(user_id, task_id, 'completed')

# ==============================================================================
# 路由逻辑
# ==============================================================================

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/verify_login', methods=['POST'])
def verify_login():
    key = request.json.get('key', '').strip()
    if is_valid_key(key):
        return jsonify({"status": "success", "msg": "登录成功"})
    else:
        return jsonify({"status": "fail", "msg": "无效的卡密"}), 401

@app.route('/control', methods=['POST'])
def control_task():
    if not check_auth(): return jsonify({"error": "无效的卡密"}), 401
    user_id = request.headers.get('X-User-ID')
    data = request.json
    task_id = data.get('task_id')
    action = data.get('action')
    if action == 'pause': task_manager.set_status(user_id, task_id, 'paused')
    elif action == 'resume': task_manager.set_status(user_id, task_id, 'running')
    elif action == 'stop': task_manager.set_status(user_id, task_id, 'stopped')
    return jsonify({"status": "success"})

@app.route('/export_docx', methods=['POST'])
def export_docx():
    if not check_auth(): return jsonify({"error": "无效的卡密"}), 401
    data = request.json
    try:
        file_stream = MarkdownToDocx.convert(data.get('content', ''))
        return send_file(file_stream, as_attachment=True, download_name='thesis.docx')
    except Exception as e:
        return jsonify({"error": str(e)}), 500

def check_auth():
    user_id = request.headers.get('X-User-ID')
    if not user_id or user_id not in VALID_KEYS: return False
    return True

@app.route('/generate', methods=['POST'])
def generate_start():
    if not check_auth(): return jsonify({"error": "Unauthorized"}), 401
    user_id = request.headers.get('X-User-ID')
    
    # 获取表单数据
    raw_chapters = request.form.get('chapter_data')
    title = request.form.get('title')
    references = request.form.get('references')
    text_custom_data = request.form.get('custom_data', '')
    task_id = request.form.get('task_id')
    initial_context = request.form.get('initial_context', '')
    
    # 读取文件流到内存
    uploaded_files = request.files.getlist('data_files')
    raw_files_data = []
    
    if uploaded_files:
        for file in uploaded_files:
            if file.filename:
                file_content = io.BytesIO(file.read())
                raw_files_data.append({
                    'name': file.filename, 
                    'content': file_content
                })

    # 初始化任务
    task_manager.start_task(user_id, task_id)
    writer = PaperAutoWriter(API_KEY, BASE_URL, MODEL_NAME)
    
    def check_status_func():
        return task_manager.get_status(user_id, task_id)

    # 启动后台线程
    t = threading.Thread(
        target=background_worker,
        args=(writer, task_id, title, json.loads(raw_chapters), references, text_custom_data, raw_files_data, check_status_func, initial_context, user_id)
    )
    t.daemon = True 
    t.start()

    return jsonify({"status": "success", "msg": "Task started in background"})

@app.route('/stream_progress')
def stream_progress():
    if not check_auth(): return "Unauthorized", 401
    
    user_id = request.headers.get('X-User-ID')
    task_id = request.args.get('task_id')
    try: last_event_index = int(request.args.get('last_index', 0))
    except: last_event_index = 0

    def event_stream():
        current_idx = last_event_index
        
        while True:
            # 获取新消息
            events, status = task_manager.get_events_from(user_id, task_id, current_idx)
            
            if events:
                for event in events:
                    event_str = str(event)
                    if not event_str.endswith('\n\n'):
                        event_str += '\n\n'
                    yield event_str
                    current_idx += 1
            else:
                if status in ['stopped', 'completed']:
                    yield f"data: {json.dumps({'type': 'done'})}\n\n"
                    break
                
                # 心跳包频率
                yield ": keep-alive\n\n"
                time.sleep(0.5) 

    # 禁用缓存
    response = Response(stream_with_context(event_stream()), content_type='text/event-stream')
    response.headers['X-Accel-Buffering'] = 'no'
    response.headers['Cache-Control'] = 'no-cache'

    
    return response

# --- 管理员相关接口 ---
@app.route('/admin')
def admin_page(): return render_template('admin.html')

@app.route('/api/admin/login', methods=['POST'])
def admin_login():
    if request.json.get('username') == ADMIN_USERNAME and request.json.get('password') == ADMIN_PASSWORD:
        session['is_admin'] = True
        return jsonify({"status": "success"})
    return jsonify({"status": "fail"}), 401

@app.route('/api/admin/logout', methods=['POST'])
def admin_logout():
    session.pop('is_admin', None)
    return jsonify({"status": "success"})

@app.route('/api/admin/keys', methods=['GET', 'POST', 'DELETE'])
def manage_keys():
    if not session.get('is_admin'): return "Unauthorized", 401
    if request.method == 'GET': return jsonify({"keys": list(VALID_KEYS)})
    if request.method == 'DELETE':
        key = request.json.get('key')
        if key in VALID_KEYS: VALID_KEYS.remove(key); save_keys(list(VALID_KEYS))
        return jsonify({"status": "success"})
    if request.method == 'POST':
        custom = request.json.get('key', '').strip()
        new_key = custom if custom else f"key_{secrets.token_hex(4)}"
        if new_key in VALID_KEYS: return jsonify({"status": "fail", "msg": "Exists"}), 400
        VALID_KEYS.add(new_key); save_keys(list(VALID_KEYS))
        return jsonify({"status": "success", "key": new_key})

# ==============================================================================
# 启动入口 (核心修改)
# ==============================================================================
if __name__ == '__main__':
    # 确保存储文件存在
    if not os.path.exists(KEYS_FILE):
        VALID_KEYS.add("test_vip_888")
        save_keys(list(VALID_KEYS))
    
    print("🚀 服务器正在启动...")
    print("⚠️  请访问 http://192.168.0.35:8001 (请根据实际IP访问)")
    print("✅ 已启用 Waitress 高并发模式，支持多任务同时运行")
    
    # ❌ 不再使用 app.run()，它不适合并发 SSE
    # app.run(debug=True, host="0.0.0.0", port=8001, threaded=True)
    
    # ✅ 使用 Waitress 启动，配置 10 个处理线程
    serve(app, host="0.0.0.0", port=8001, threads=100, connection_limit=200, channel_timeout=300)