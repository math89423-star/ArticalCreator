import pandas as pd
import pypdf
import docx
import base64
import io

def extract_file_content(file_stream, filename, llm_client=None) -> str:
    """
    根据文件后缀名，提取文件内容为纯文本字符串。
    [新增] 支持图片解析 (需要传入 llm_client)
    """
    filename = filename.lower()
    raw_text = "" 
    
    try:
        # 重置指针
        if hasattr(file_stream, 'seek'):
            file_stream.seek(0)

        if filename.endswith(('.png', '.jpg', '.jpeg', '.bmp', '.webp')):
            if not llm_client:
                return "【系统提示】解析图片需要 LLM 客户端支持，当前未提供。"
            
            try:
                # 1. 读取图片并转 Base64
                img_bytes = file_stream.read()
                b64_str = base64.b64encode(img_bytes).decode('utf-8')
                mime_type = "image/png" if filename.endswith('.png') else "image/jpeg"
                
                # 2. 调用 Vision 模型进行“读图”
                # Prompt 设计：要求模型详细描述图片中的数据、趋势和文字
                response = llm_client.chat.completions.create(
                    model="gemini-2.5-pro", # 确保使用支持 Vision 的模型
                    messages=[
                        {
                            "role": "system", 
                            "content": "你是一个数据分析师。请仔细观察这张图片，提取其中所有的关键信息、数据趋势、文字内容，并整理成结构化的文本描述，以便后续用于论文写作。"
                        },
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": f"请详细描述这张名为 '{filename}' 的图片内容："},
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:{mime_type};base64,{b64_str}"
                                    }
                                }
                            ]
                        }
                    ],
                    max_tokens=2000
                )
                raw_text = f"[图片视觉解析结果]:\n{response.choices[0].message.content}"
            
            except Exception as e:
                print(f"图片解析失败: {e}")
                raw_text = f"图片解析失败: {str(e)}"

        # ==========================================
        # 2. 常规文本文件处理 (保持不变)
        # ==========================================
        # 1. Excel/CSV
        elif filename.endswith('.csv'):
            try:
                df = pd.read_csv(file_stream)
            except UnicodeDecodeError:
                file_stream.seek(0)
                df = pd.read_csv(file_stream, encoding='gbk')
            raw_text = df.head(60).to_markdown(index=False)
        
        elif filename.endswith(('.xls', '.xlsx')):
            df = pd.read_excel(file_stream)
            raw_text = df.head(60).to_markdown(index=False)
            
        # 2. TXT
        elif filename.endswith('.txt'):
            try:
                raw_text = file_stream.read().decode('utf-8')
            except:
                file_stream.seek(0)
                raw_text = file_stream.read().decode('gbk', errors='ignore')
            raw_text = raw_text[:8000] 
            
        # 3. PDF
        elif filename.endswith('.pdf'):
            reader = pypdf.PdfReader(file_stream)
            for i, page in enumerate(reader.pages[:15]): 
                page_text = page.extract_text()
                if page_text: 
                    raw_text += f"[Page {i+1}] {page_text}\n"

        # 4. DOCX
        elif filename.endswith('.docx'):
            doc = docx.Document(file_stream)
            for para in doc.paragraphs:
                if para.text.strip():
                    raw_text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    raw_text += " | ".join(row_text) + "\n"
            raw_text = raw_text[:8000]

        else:
            return "" 
            
    except Exception as e:
        print(f"解析文件 {filename} 失败: {e}")
        return "" 

    # 返回 XML 包裹的结构化数据
    return f"""
<datasource name="{filename}">
{raw_text}
</datasource>
"""